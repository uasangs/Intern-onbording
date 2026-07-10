
import os
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image,
    Table, TableStyle, HRFlowable, ListFlowable, ListItem
)

from app.core.config import settings
import subprocess

UPLOAD_DIR    = settings.UPLOAD_DIR
TEMPLATES_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "..", "pdf", "templates"
)
LOGO_PATH      = os.path.join(TEMPLATES_DIR, "grasim_logo.jpg")
SIGNATURE_PATH = os.path.join(TEMPLATES_DIR, "signature.png")

# ── Page geometry (from docx sectPr twips ÷ 20 = pt) ─────────────────────────
PAGE_W, PAGE_H = A4
ML   = 78.0    # left  1560 twips
MR   = 71.0    # right 1419 twips
MT   = 105.75  # top   2115 twips
MB   = 105.0   
HEADER_DIST = 13.7   # 274 twips
FOOTER_DIST = 82.75  # 1655 twips

# Logo exact EMU from docx header
LOGO_W  = (922655 / 914400) * 72   # 72.65 pt
LOGO_H  = (906145 / 914400) * 72   # 71.35 pt
LOGO_CX = PAGE_W / 2

# Signature exact EMU from docx body
SIG_W = (1444530 / 914400) * 72    # 113.9 pt
SIG_H = (800100  / 914400) * 72    # 63.1  pt

# ── Fonts ─────────────────────────────────────────────────────────────────────
FONT      = "Helvetica"       # Arial → Helvetica (built-in PDF equivalent)
FONT_BOLD = "Helvetica-Bold"
FONT_BI   = "Helvetica-BoldOblique"  # Bold+Italic for "Do note"
FONT_I    = "Helvetica-Oblique"
FONT_HDR  = "Helvetica-Bold"  # Tahoma → Helvetica-Bold for header label
FS        = 11    # PlainText style body
FS_NORM   = 12    # Normal style (not used much)
FS_SM     = 8     # footer small
FS_CO     = 10    # footer company name first line
BLACK     = HexColor("#000000")


def _s(name, **kw):
    base = dict(
        fontName=FONT, fontSize=FS, leading=FS * 1.5,
        textColor=BLACK, spaceBefore=0, spaceAfter=0,
    )
    base.update(kw)
    return ParagraphStyle(name, **base)


def _styles():
    return dict(
        n    = _s("n",    alignment=TA_JUSTIFY),
        c    = _s("c",    alignment=TA_CENTER),
        l    = _s("l",    alignment=TA_LEFT),
        blt  = _s("blt",  alignment=TA_JUSTIFY,
                  leftIndent=22.5, firstLineIndent=-22.5),
        sml  = _s("sml",  fontSize=FS_SM,  leading=FS_SM*1.4,  alignment=TA_LEFT),
        smc  = _s("smc",  fontSize=FS_SM,  leading=FS_SM*1.4,  alignment=TA_CENTER),
        smco = _s("smco", fontSize=FS_CO,  leading=FS_CO*1.4,  alignment=TA_CENTER),
    )


def _ensure_dir(p): os.makedirs(p, exist_ok=True)

def _ordinal(n):
    sfx = "th" if 10 <= n % 100 <= 20 else {1:"st",2:"nd",3:"rd"}.get(n%10,"th")
    return f"{n}{sfx}"

def _fmt_long(d):
    return f"{_ordinal(d.day)} {d.strftime('%B %Y')}" if d else ""

def _fmt_letter(d):
    return d.strftime("%d-%b-%Y") if d else ""

def _get_hr_signatory(record):
    hr_name = settings.HR_HEAD_NAME
    hr_div  = record.location or ""
    try:
        from app.core.database import SessionLocal
        from app.models.models import MasterData
        db = SessionLocal()
        try:
            row = db.query(MasterData).filter(MasterData.id == 1).first()
            if row and row.letter_formats:
                fmt = next((f for f in row.letter_formats
                            if f.get("department") == record.department), None)
                if fmt:
                    hr_name = fmt.get("signatory", hr_name).split("\n")[0].strip()
                    hr_div  = fmt.get("department", hr_div)
        finally:
            db.close()
    except Exception:
        pass
    return hr_name, hr_div


# ── Header / Footer callback ───────────────────────────────────────────────────
def _hf_callback(logo_path, division):
    def draw(c, doc):
        c.saveState()

        # ── HEADER ────────────────────────────────────────────────────────────
        header_zone_h = MT - HEADER_DIST
        logo_y = PAGE_H - HEADER_DIST - (header_zone_h + LOGO_H) / 2 + 4

        if os.path.exists(logo_path):
            c.drawImage(
                logo_path, LOGO_CX - LOGO_W / 2, logo_y,
                width=LOGO_W, height=LOGO_H,
                preserveAspectRatio=True, mask="auto"
            )

        # "PULP & FIBRE BUSINESS" — Tahoma Bold 8pt centred
        c.setFont(FONT_HDR, 8)
        c.setFillColor(BLACK)
        c.drawCentredString(PAGE_W / 2, logo_y - 11, division.upper())

        # Bottom border on header (thin line at body top)
        # c.setStrokeColor(BLACK)
        # c.setLineWidth(0.5)
        # c.line(ML, PAGE_H - MT + 4, PAGE_W - MR, PAGE_H - MT + 4)

        # ── FOOTER ────────────────────────────────────────────────────────────
        # Footer zone starts at MB from bottom
        # Structure (top to bottom):
        #   thin rule
        #   [Birla Cellulose logo - EMF, we skip/approximate]
        #   "Grasim Industries Limited (Pulp & Fibre Business)" 10pt orange/dark
        #   address lines 8pt
        #   "Sensitivity: General" very bottom

        footer_top = MB - 2
        # c.setStrokeColor(BLACK)
        # c.setLineWidth(0.5)
        # c.line(ML, footer_top, PAGE_W - MR, footer_top)

        # Birla Cellulose logo — centred above company name
        BIRLA_LOGO = os.path.join(TEMPLATES_DIR, "birla_cellulose_logo.png")
        BIRLA_W = 50.0   # pt  (~0.69 inch, matching docx footer visual)
        BIRLA_H = 54.0   # pt  (aspect ratio of cropped logo)
        if os.path.exists(BIRLA_LOGO):
            c.drawImage(
                BIRLA_LOGO,
                LOGO_CX - BIRLA_W / 2,
                footer_top - 4 - BIRLA_H,
                width=BIRLA_W, height=BIRLA_H,
                preserveAspectRatio=True, mask="auto"
            )
            co_base = footer_top - 4 - BIRLA_H - 12
        else:
            co_base = footer_top - 13

        # Company name in reddish-brown color (from docx footer color)
        FOOTER_CO_COLOR = HexColor("#C0392B")  # dark red matching Birla Cellulose branding
        c.setFont(FONT_BOLD, FS_CO)
        c.setFillColor(FOOTER_CO_COLOR)
        co_y = co_base
        c.drawCentredString(PAGE_W / 2, co_y,
                            "Grasim Industries Limited (Pulp & Fibre Business)")

        # Address lines
        c.setFont(FONT, FS_SM)
        c.setFillColor(BLACK)
        lines = [
            "Unit 501/A, 502, 5th Floor, Hubtown Solaris, Prof. N.S. Phadke Marg,"
            " Vijay Nagar, Andheri (E), Mumbai 400069, India",
            "T: + 91 2261957700  |  F: + 91 2261957702  |"
            "  E: birlacellulose@adityabirla.com  |  W: www.birlacellulose.com",
            f"Corporate ID No. : {settings.COMPANY_CIN}",
        ]
        line_y = co_y - 10
        for line in lines:
            c.drawCentredString(PAGE_W / 2, line_y, line)
            line_y -= 9

        # "Sensitivity: General" at very bottom left
        c.setFont(FONT, 7.5)
        c.setFillColor(BLACK)
        c.drawString(ML, line_y - 4, "Sensitivity: General")

        c.restoreState()
    return draw


# ── Offer Letter ──────────────────────────────────────────────────────────────
async def generate_offer_letter_pdf(record) -> str:
    cand = record.candidate

    full_name  = cand.full_name if cand else "Candidate"
    first_name = full_name.split()[0] if full_name else "Candidate"
    gender     = (cand.gender if cand else "male") or "male"
    sal        = "Mr." if gender.lower() == "male" else "Ms."

    institute  = cand.institute_name or ""
    city       = cand.city or ""
    state      = cand.state or ""
    location   = record.location or ""
    weeks      = record.duration_weeks or ""

    ldate   = _fmt_letter(record.start_date or datetime.now())
    start_s = _fmt_long(record.start_date)
    end_s   = _fmt_long(record.end_date)

    has_stip = record.stipend_amount and float(record.stipend_amount) > 0
    stip_s   = f"Rs {int(record.stipend_amount):,}/- per month" if has_stip else ""

    hr_name, hr_div = _get_hr_signatory(record)
    co = settings.COMPANY_NAME

    import io
    S  = _styles()
    cb = _hf_callback(LOGO_PATH, hr_div or settings.COMPANY_DIVISION)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=ML, rightMargin=MR,
        topMargin=MT, bottomMargin=MB,
        title=f"Offer Letter – {full_name}",
    )

    def sp(h=12): return Spacer(1, h)

    story = []

    # Date
    story.append(Paragraph(ldate, S["n"]))
    story.append(sp(12))

    # Addressee block
    story.append(Paragraph(f"{sal} {full_name}", S["n"]))
    story.append(Paragraph(institute, S["n"]))
    addr = city
    if state:
        addr += f"<br/>District- {state}"
    story.append(Paragraph(addr, S["n"]))
    story.append(sp(12))

    # Subject — bold
    story.append(Paragraph(
        f"<b>Sub: Regarding your application for Internship with {co}</b>",
        S["n"]))
    story.append(sp(12))

    # Salutation
    story.append(Paragraph(f"Dear {first_name},", S["n"]))
    story.append(sp(12))

    # Body paragraph 1
    story.append(Paragraph(
        f"This is in regard to your request letter for internship with {co} at "
        f"<b>{location}</b>. We are pleased to inform you that you have been accepted "
        f"as an Intern for a period of {weeks} weeks starting from "
        f"<b>{start_s}</b> to <b>{end_s}</b>.",
        S["n"]))
    story.append(sp(12))

    # Body paragraph 2
    story.append(Paragraph(
        f"During the Internship you shall be bound by all terms and conditions "
        f"attached in Annexure A and Annexure B with this letter. You shall further "
        f"be bound by all such internal policies of {co} and such changes as may be "
        f"informed to you from time to time.",
        S["n"]))
    story.append(sp(12))

    # "Do note" — Bold + Italic + Underline
    story.append(Paragraph(
        f"<b><i><u>Do note that this is not an employment offer.</u></i></b>",
        S["n"]))
    story.append(sp(12))

    story.append(Paragraph(
        "The following additional terms &amp; conditions will be applicable to you "
        "during the Internship:",
        S["n"]))
    story.append(sp(6))

    # Bullet list — large bullet, hanging indent matching docx
    bullets = [
        "You will Work from office for the Project assigned to you.",
        "You may have to Visit the customers / vendors, based on the requirement of the Project.",
    ]
    if has_stip:
        bullets.append(
            f"You will be paid a Stipend of <b>{stip_s}</b>, inclusive of all taxes.")
    bullets += [
        f"If there is a requirement for you to travel to complete your Project, {co} "
        f"will reimburse the travel expenses as per the original bills and in accordance "
        f"with the travel policy of {co}. Any travel during the term of your Internship "
        f"shall be arranged for by you. {co} shall not be held responsible in case of "
        f"any untoward event / incident during travel.",
        "You shall arrange for your own accommodation during the Internship tenure.",
        f"The certificate for the project completion will be issued on the submission of "
        f"the project report to the undersigned. Please take note that this internship is "
        f"permitted only as it is a part of the curriculum. After the completion of "
        f"internship, {co} is not liable to offer the student any type of employment.",
        "This engagement is subject to you signing Annexure A and Annexure B attached "
        "with this letter and returning a copy of the same to us.",
    ]

    blt_style = _s("blt2", alignment=TA_JUSTIFY,
                   leftIndent=22.5, firstLineIndent=0, spaceAfter=0)
    for b in bullets:
        story.append(Paragraph(f"\u2022\u00a0\u00a0{b}", blt_style))

    story.append(sp(12))

    story.append(Paragraph(
        "Please countersign the duplicate copy of this offer letter and return to us "
        "as a token of your acceptance.",
        S["n"]))
    story.append(sp(12))

    story.append(Paragraph("Yours faithfully,", S["n"]))
    story.append(sp(6))

    # Signature image
    if os.path.exists(SIGNATURE_PATH):
        sig = Image(SIGNATURE_PATH, width=SIG_W, height=SIG_H)
        sig.hAlign = "LEFT"
        story.append(sig)
    else:
        story.append(sp(SIG_H))

    # Signatory — bold
    story.append(Paragraph(f"<b>{hr_name}</b>", S["n"]))
    story.append(Paragraph(f"<b>{settings.HR_HEAD_TITLE}</b>", S["n"]))
    story.append(Paragraph(f"<b>{hr_div or settings.COMPANY_DIVISION}</b>", S["n"]))

    story.append(sp(24))  # blank lines before Encl

    # Enclosures — plain, matching docx
    story.append(Paragraph(
        "Encl: (i)  Annexure A: General terms and conditions", S["n"]))
    story.append(Paragraph(
        "         (ii) Annexure B: Declaration and indemnity", S["n"]))

    story.append(sp(24))  # blank lines before Declaration

    # DECLARATION — centered, no HR line (docx uses blank paragraphs only)
    story.append(Paragraph("DECLARATION", S["c"]))
    story.append(sp(12))

    story.append(Paragraph(
        "I hereby declare and affirm that I have carefully studied and understood the "
        "terms and conditions of my internship herein above detailed and specifically "
        "mentioned in Annexure A and Annexure B. I accept the said internship and I "
        "hereby undertake to abide myself of the said terms and conditions.",
        S["n"]))

    # Declaration fields — each on its own line with underline space, matching docx
    story.append(sp(12))
    story.append(Paragraph("Place:", S["n"]))
    story.append(sp(12))
    story.append(Paragraph("Date:", S["n"]))
    story.append(sp(12))

    # Name of University + Signature on same line (tab-spaced in docx)
    sig_line_tbl = Table(
    [["Name of the University:", "", "Signature", ""]],
    colWidths=[1.8*72, 1.5*72, 0.9*72, 1.5*72]
)
    # sig_line_tbl.setStyle(TableStyle([
    #     ("FONTNAME",     (0,0), (-1,-1), FONT),
    #     ("FONTSIZE",     (0,0), (-1,-1), FS),
    #     ("VALIGN",       (0,0), (-1,-1), "BOTTOM"),
    #     ("TOPPADDING",   (0,0), (-1,-1), 0),
    #     ("BOTTOMPADDING",(0,0), (-1,-1), 0),
    #     ("LEFTPADDING",  (0,0), (-1,-1), 0),
    #     ("RIGHTPADDING", (0,0), (-1,-1), 2),
    # ]))
    story.append(sig_line_tbl)

    doc.build(story, onFirstPage=cb, onLaterPages=cb)
    return buffer.getvalue()  # Return bytes — stored in DB by caller

# ── Experience Certificate ────────────────────────────────────────────────────

def _cert_ordinal(n):
    sfx = "th" if 10 <= n % 100 <= 20 else {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{sfx}"

def _cert_fmt_long(d):
    return f"{_cert_ordinal(d.day)} {d.strftime('%B %Y')}" if d else ""

def _cert_clear_and_set(paragraph, new_text):
    """Replace all runs with a single run, preserving first run formatting."""
    from docx.oxml.ns import qn
    first_run = next((r for r in paragraph.runs if r.text.strip()), None)
    p_elem = paragraph._element
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    run = paragraph.add_run(new_text)
    if first_run:
        run.bold = first_run.bold
        run.italic = first_run.italic
        run.underline = first_run.underline
        if first_run.font.name: run.font.name = first_run.font.name
        if first_run.font.size: run.font.size = first_run.font.size

def _cert_set_runs(paragraph, runs_data):
    """Set paragraph with multiple runs. runs_data: list of (text, bold, italic, underline)"""
    from docx.oxml.ns import qn
    first_run = paragraph.runs[0] if paragraph.runs else None
    font_name = first_run.font.name if first_run else None
    font_size = first_run.font.size if first_run else None
    p_elem = paragraph._element
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    for text, bold, italic, underline in runs_data:
        run = paragraph.add_run(text)
        run.bold = bold; run.italic = italic; run.underline = underline
        if font_name: run.font.name = font_name
        if font_size: run.font.size = font_size

async def generate_certificate_pdf(record, payload) -> bytes:
    """
    Generate experience certificate PDF using the official docx template.
    Substitutes text only — all images (logo, signature) preserved from template.
    Uses docx2pdf (MS Word on Windows) for pixel-perfect PDF conversion.
    """
    import io, tempfile, os
    from docx import Document

    template_path = os.path.join(TEMPLATES_DIR, "certificate_template.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Certificate template not found: {template_path}")

    doc = Document(template_path)

    cand = record.candidate
    full_name   = cand.full_name if cand else "Candidate"
    gender      = (cand.gender if cand else "male") or "male"
    sal         = "Mr." if gender.lower() == "male" else "Ms."
    pronoun     = "his" if gender.lower() == "male" else "her"
    pronoun_obj = "him" if gender.lower() == "male" else "her"
    institute   = cand.institute_name or ""
    location    = record.location or ""
    weeks       = record.duration_weeks or ""
    conduct     = payload.conduct_remark.lower()
    project     = payload.project_title
    guides      = payload.guide_names
    start_s     = _cert_fmt_long(record.start_date)
    end_s       = _cert_fmt_long(record.end_date)
    issue_date  = payload.issue_date
    issue_s = f"{_cert_ordinal(issue_date.day)} {issue_date.strftime('%B')} {issue_date.year}"
    hr_name, hr_div = _get_hr_signatory(record)

    MONTHS = ["January","February","March","April","May","June",
              "July","August","September","October","November","December"]

    for para in doc.paragraphs:
        text = para.text.strip()

        # Date line
        if any(m in text for m in MONTHS) and len(text) < 25 and any(c.isdigit() for c in text):
            _cert_clear_and_set(para, issue_s)

        # Certify paragraph
        elif "This is to certify that" in text:
            _cert_set_runs(para, [
                ("This is to certify that ", False, False, False),
                (f"{sal} {full_name}, ", True, False, False),
                (f"a student of {institute} has successfully completed "
                 f"{pronoun} Internship at {hr_div or location} (Grasim Industries Limited), "
                 f"for the period of {weeks} weeks starting from {start_s} to {end_s}.",
                 False, False, False),
            ])

        # Conduct paragraph
        elif "During" in text and "tenure" in text and "conduct" in text:
            _cert_set_runs(para, [
                (f"During {pronoun} tenure with us as Intern {pronoun} conduct was ", False, False, False),
                (conduct + ".", True, False, False),
            ])

        # Project paragraph
        elif "project undertaken" in text:
            _cert_set_runs(para, [
                ("The project undertaken was ", False, False, False),
                (f'"{project}" ', True, False, False),
                ("under the guidance of ", False, False, False),
                (guides + ".", True, False, False),
            ])

        # Closing — signature image is embedded in this paragraph, preserve image runs
        elif "We wish" in text and "future endeavor" in text:
            from docx.oxml.ns import qn
            p_elem = para._element
            # Only remove text runs, keep drawing/image runs
            for r in p_elem.findall(qn('w:r')):
                # Check if this run contains a drawing (image) — if so, keep it
                if r.find(qn('w:drawing')) is None:
                    p_elem.remove(r)
            # Add text run at the beginning
            import copy
            new_run = para.add_run(
                f"We wish {pronoun_obj} all the best in {pronoun} future endeavor. "
                f"For Grasim Industries Ltd."
            )
            # Move the new run to before the image run
            image_runs = [r for r in p_elem.findall(qn('w:r'))
                         if r.find(qn('w:drawing')) is not None]
            text_runs = [r for r in p_elem.findall(qn('w:r'))
                        if r.find(qn('w:drawing')) is None]
            if image_runs and text_runs:
                # Insert text run before image run
                p_elem.remove(text_runs[0])
                p_elem.insert(list(p_elem).index(image_runs[0]), text_runs[0])

        # Signatory name
        elif text == "Sheba Banerjee":
            _cert_clear_and_set(para, hr_name)

        # Head HR + division
        elif "Head" in text and "Human Resources" in text:
            _cert_clear_and_set(para, f"Head \u2013 Human Resources {hr_div}")

    # Save docx to temp file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f)
        docx_path = f.name

    pdf_path = docx_path.replace('.docx', '.pdf')

    try:
        import platform
        if platform.system() == 'Windows':
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        else:
            env = os.environ.copy()
            env['HOME'] = '/tmp/libreoffice-home'
            os.makedirs('/tmp/libreoffice-home', exist_ok=True)
            result = subprocess.run([
               'libreoffice', '--headless', '--convert-to', 'pdf',
               '--outdir', os.path.dirname(pdf_path), docx_path
            ], capture_output=True, timeout=60, env=env)
        if os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                return f.read()
    finally:
        if os.path.exists(docx_path):
            os.unlink(docx_path)
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)

    # Fallback: return docx bytes (if MS Word not available)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()